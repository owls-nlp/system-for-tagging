import os
import sys
import time
import random
import string
import re
import math
import subprocess
import copy

import pandas as pd
import numpy as np
import cv2
from PIL import Image
import pytesseract
import pdfplumber

import glob
import multiprocessing as mp
import tqdm

import detectron2
from detectron2.config import get_cfg
from detectron2.data.detection_utils import read_image
from detectron2.utils.logger import setup_logger

from predictor import VisualizationDemo

from detectron2.data import MetadataCatalog

from docx import Document as DocumentDocx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

import torch.nn as nn
import torchvision.transforms.functional as TF
from ssd import build_ssd
import torch
from torch.autograd import Variable

class FindingFormulasModel:
    def __init__(self, path_to_model: str, score_thresh_test: float):
        exp_cfg = {
                'hboxes512': {
                'num_classes': 2,
                'lr_steps': (80000, 100000, 120000),
                'max_iter': 132000,
                'feature_maps': [64, 32, 16, 8, 4, 2, 1],
                'min_dim': 512,
                'steps': [8, 16, 32, 64, 128, 256, 512],
                'min_sizes': [35.84, 76.8, 153.6, 230.4, 307.2, 384.0, 460.8],
                'max_sizes': [76.8, 153.6, 230.4, 307.2, 384.0, 460.8, 537.6],
                'aspect_ratios': [[2, 3, 5, 7, 10], [2, 3, 5, 7, 10], [2, 3, 5, 7, 10], [2, 3, 5, 7, 10],
                                [2, 3, 5, 7, 10], [2, 3, 5, 7, 10], [2, 3, 5, 7, 10]],
                'variance': [0.1, 0.2],
                'clip': True,
                'name': 'ssd512',
                'is_vertical_prior_boxes_enabled': False,
                'mbox': {
                    '512': [7,7,7,7,7,7,7],
                },
                'extras': {
                    '512': [256, 'S', 512, 128, 'S', 256, 128, 'S', 256, 128, 'S', 256],
                }
            }
        }

        self.net = build_ssd(exp_cfg['hboxes512'])
        self.net.to('cpu')
        self.net = nn.DataParallel(self.net)
        self.net.load_state_dict(torch.load(path_to_model, map_location=torch.device('cpu')))
        self.net.eval()
        self.score_thresh_test = score_thresh_test
    

    def predict(self, img_path):
        image = cv2.imread(img_path)
        rgb_image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
        image = cv2.resize(image, (512, 512)).astype(np.float32)
        image = image[:, :, ::-1].copy()
        image = torch.from_numpy(image).permute(2, 0, 1)
        
        images = Variable(image.unsqueeze(0))

        y, debug_boxes, debug_scores = self.net(images)  # forward pass
        detections = y.data

        result_dict = dict()

        scale = torch.Tensor(rgb_image.shape[1::-1]).repeat(2)
        n_i = 0
        for i in range(detections.size(1)):
            j = 0
            while detections[0,i,j,0] > self.score_thresh_test:
                score = detections[0,i,j,0]
                pt = (detections[0,i,j,1:]*scale).cpu().numpy()
                coords = [pt[0], pt[1], pt[2], pt[3]]
                result_dict[n_i] = coords
                n_i += 1
                j+=1
        
        df = pd.DataFrame.from_dict(data=result_dict, orient='index', columns = ['left', 'upper', 'right', 'lower'])
        if df.shape[0] > 0:
            delta = 300
            df.left = df.left - delta
            df.upper = df.upper - delta
            df.right = df.right + delta
            df.lower = df.lower + delta
        
        return df


class SegmentationModel:
    '''Class for run detectron2 model'''
    def __init__(self, path_to_model: str, path_to_cfg_config: str, device: str, score_thresh_test: float):
        self.cfg = get_cfg()
        self.cfg.merge_from_file(path_to_cfg_config)
        self.cfg.MODEL.ROI_HEADS.SCORE_THRESH_TEST = score_thresh_test
        self.cfg.MODEL.WEIGHTS = path_to_model
        self.cfg.MODEL.DEVICE = device
        self.classes = ['text', 'title', 'list', 'table', 'figure']
        self.default_predictor = detectron2.engine.defaults.DefaultPredictor(self.cfg)

        # remove it!!! #TODO
        # from predictor import VisualizationDemo
        # self.demo = VisualizationDemo(self.cfg)
    

    def calc_iou(self, bb1: dict, bb2: dict):
        """
        Calculate the Intersection over Union (IoU) of two bounding boxes.
        Parameters
        ----------
        bb1 : dict
            Keys: {'x1', 'x2', 'y1', 'y2'}
            The (x1, y1) position is at the top left corner,
            the (x2, y2) position is at the bottom right corner
        bb2 : dict
            Keys: {'x1', 'x2', 'y1', 'y2'}
            The (x, y) position is at the top left corner,
            the (x2, y2) position is at the bottom right corner
        Returns
        -------
        float
            in [0, 1]
        """
        assert bb1['x1'] < bb1['x2']
        assert bb1['y1'] < bb1['y2']
        assert bb2['x1'] < bb2['x2']
        assert bb2['y1'] < bb2['y2']

        # determine the coordinates of the intersection rectangle
        x_left = max(bb1['x1'], bb2['x1'])
        y_top = max(bb1['y1'], bb2['y1'])
        x_right = min(bb1['x2'], bb2['x2'])
        y_bottom = min(bb1['y2'], bb2['y2'])

        if x_right < x_left or y_bottom < y_top:
            return 0.0, 0.0, 0.0

        # The intersection of two axis-aligned bounding boxes is always an
        # axis-aligned bounding box
        intersection_area = (x_right - x_left) * (y_bottom - y_top)

        # compute the area of both AABBs
        bb1_area = (bb1['x2'] - bb1['x1']) * (bb1['y2'] - bb1['y1'])
        bb2_area = (bb2['x2'] - bb2['x1']) * (bb2['y2'] - bb2['y1'])

        # compute the intersection over union by taking the intersection
        # area and dividing it by the sum of prediction + ground-truth
        # areas - the interesection area
        iou = intersection_area / float(bb1_area + bb2_area - intersection_area)
        if iou < 0:
            return -1.0, 0.0, 0.0
        elif iou > 1:
            return -1.0, 0.0, 0.0
        else:
            return iou, intersection_area / bb1_area, intersection_area / bb2_area
    
    
    def predict(self, img, layout_type):
        predictions = self.default_predictor(img)
        #_, visualized_output = self.demo.run_on_image(img)
        #visualized_output.save('output_all.jpg')
        instances = predictions["instances"].to('cpu')
        pred_classes = instances.pred_classes
        labels = [self.classes[i] for i in pred_classes]
        boxes = instances.pred_boxes
        scores = instances.scores

        if isinstance(boxes, detectron2.structures.boxes.Boxes):
            boxes = boxes.tensor.numpy()
        else:
            boxes = np.asarray(boxes)

        img = Image.fromarray(img)

        df = pd.DataFrame(boxes, columns=['left', 'upper', 'right', 'lower'])
        df['label'] = labels
        df['scores'] = scores
        df['image_width'] = img.size[0]

        df['left_round'] = df.left.apply(lambda x: round(x/10)*10)
        df['upper_round'] = df.upper.apply(lambda x: round(x/10)*10)
        df['right_round'] = df.right.apply(lambda x: round(x/10)*10)
        df['lower_round'] = df.lower.apply(lambda x: round(x/10)*10)

        if layout_type == 0:

            df = df.sort_values(by=['upper_round', 'left_round'], ascending=True)
            df = df.reset_index(drop=True)

            df['type'] = df.apply(lambda x: 'one' if (x['right']-x['left']) >= x['image_width']/2.0 else 'two', axis=1)

            flag_one = False
            flag_two = False
            i_df = 0
            i_dict = 0
            tmp_list = list()
            result_dict = dict()
            while i_df < df.shape[0]:
                if df.type.iloc[i_df] == 'one' and flag_one == False and flag_two == True:
                    flag_one = True
                    flag_two = False
                    result_dict[i_dict] = tmp_list
                    tmp_list = list()
                    tmp_list.append(i_df)
                    i_dict += 1
                elif df.type.iloc[i_df] == 'two' and flag_one == True and flag_two == False:
                    flag_one = False
                    flag_two = True
                    result_dict[i_dict] = tmp_list
                    tmp_list = list()
                    tmp_list.append(i_df)
                    i_dict += 1
                elif df.type.iloc[i_df] == 'two' and flag_two == False:
                    flag_two = True
                    tmp_list.append(i_df)
                elif df.type.iloc[i_df] == 'two' and flag_two == True:
                    tmp_list.append(i_df)
                elif df.type.iloc[i_df] == 'one' and flag_one == False:
                    flag_one = True
                    tmp_list.append(i_df)
                elif df.type.iloc[i_df] == 'one' and flag_one == True:
                    tmp_list.append(i_df)
                i_df += 1

            result_dict[i_dict] = tmp_list

            result_df = pd.DataFrame()
            for i in range(len(result_dict)):
                tmp_df = df.iloc[result_dict[i], :]
                if 'two' == tmp_df['type'].iloc[0]:
                    tmp_df['marker'] = tmp_df.apply(lambda x: 'left' if x['left'] <= x['image_width']/2. else 'right', axis = 1)
                    tmp_df['right'] = tmp_df.right.apply(math.ceil)
                    
                    if tmp_df[tmp_df.marker == 'left'].shape[0] > 0:
                        min_l_left = math.floor(tmp_df[tmp_df.marker == 'left'].left.min())
                        #max_r_left = tmp_df[tmp_df.marker == 'left'].right.max()
                    
                    if tmp_df[tmp_df.marker == 'right'].shape[0] > 0:
                        min_l_right = math.floor(tmp_df[tmp_df.marker == 'right'].left.min())
                        #max_r_right = tmp_df[tmp_df.marker == 'right'].right.max()

                    result_tmp_df = pd.DataFrame()
                    
                    left_tmp_df = tmp_df[tmp_df.marker == 'left'].sort_values(by=['upper'], ascending=True)
                    if left_tmp_df.shape[0] > 0:
                        left_tmp_df.left = min_l_left
                    #left_tmp_df.right = max_r_left
                    
                    result_tmp_df = result_tmp_df.append(
                        left_tmp_df, ignore_index=True
                    )
                    
                    right_tmp_df = tmp_df[tmp_df.marker == 'right'].sort_values(by=['upper'], ascending=True)
                    if right_tmp_df.shape[0] > 0:
                        right_tmp_df.left = min_l_right
                    #right_tmp_df.right = max_r_right

                    result_tmp_df = result_tmp_df.append(
                        right_tmp_df, ignore_index=True
                    )
                    
                    tmp_df = result_tmp_df
                    
                result_df = result_df.append(tmp_df, ignore_index=True)
            
            #result_df.to_excel('show.xlsx')

            if 'marker' not in result_df.columns.tolist():
                result_df['marker'] = None

            result_df.marker = result_df.marker.fillna('one')
            
            bad_indexes = list()
            for i in range(result_df.shape[0]):
                for j in range(result_df.shape[0]):
                    bb1 = {
                        'x1': result_df.left.iloc[i],
                        'x2': result_df.right.iloc[i],
                        'y1': result_df.upper.iloc[i],
                        'y2': result_df.lower.iloc[i],
                        'marker': result_df.marker.iloc[i],
                        'label': result_df.label.iloc[i]
                    }
                    bb2 = {
                        'x1': result_df.left.iloc[j],
                        'x2': result_df.right.iloc[j],
                        'y1': result_df.upper.iloc[j],
                        'y2': result_df.lower.iloc[j],
                        'marker': result_df.marker.iloc[j],
                        'label': result_df.label.iloc[j]
                    }

                    # изменил условия на y1 и y2, + делаем проверку на маркеры

                    if bb1['y1'] >= bb2['y1'] and bb1['y2'] <= bb2['y2'] and (i!=j) and (bb1['label'] in ['text', 'title', 'list']) and (bb2['label'] in ['text', 'title', 'list']) and (bb1['marker'] == bb2['marker']):
                        bad_indexes.append(i)
                    elif (bb1['y2'] > bb2['y1']) and (i+1 == j) and (bb1['marker'] == bb2['marker']):
                        result_df.upper.iloc[j] = result_df.upper.iloc[i]
                        bad_indexes.append(i)

                    # iou, area_bb1, area_bb2 = self.calc_iou(bb1, bb2)
                    # if (iou > 0.0) and (i!=j) and (area_bb1 == 1.0):
                    #     bad_indexes.append(i)
            #print('bad', bad_indexes)

            result_df = result_df.drop(bad_indexes)
            
            for i in range(result_df.shape[0]-1):
                if result_df['type'].iloc[i] == 'one':
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]
                elif (result_df.marker.iloc[i] == 'left') and (result_df.marker.iloc[i+1] == 'left'):
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]
                elif ((result_df.marker.iloc[i] == 'right') and (result_df.marker.iloc[i+1] == 'right')) or ((result_df.marker.iloc[i] == 'right') and (result_df.marker.iloc[i+1] == 'one')):
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]
                # случай когда нет right вместе с left:
                elif (result_df.marker.iloc[i] == 'left') and (result_df.marker.iloc[i+1] == 'one'):
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]
                # самый сложный слуачй с left: поиск one после блока right или понимание это это самый нижний левый
                elif (result_df.marker.iloc[i] == 'left') and (result_df.marker.iloc[i+1] == 'right'):
                    k = copy.copy(i) + 1
                    while k <= result_df.shape[0]-1:
                        if result_df.marker.iloc[k] == 'one':
                            break
                        k += 1
                
                    # проверяем какая ситуация, если k < кол-ва данных, то значит нашли 
                    if k <= result_df.shape[0]-1:
                        result_df.lower.iloc[i] = result_df.upper.iloc[k]
                    else:
                        result_df.lower.iloc[i] += min((result_df.lower.iloc[i] - result_df.upper.iloc[i]) * 0.3, img.size[1] - result_df.lower.iloc[i])

            # прибавляем плюсом посленему элементу
            result_df.lower.iloc[result_df.shape[0]-1] += min((result_df.lower.iloc[result_df.shape[0]-1] - result_df.upper.iloc[result_df.shape[0]-1]) * 0.3, img.size[1] - result_df.lower.iloc[result_df.shape[0]-1])

            return result_df

        elif layout_type == 1:
            df = df.sort_values(by=['upper_round', 'left_round'], ascending=True)
            df = df.reset_index(drop=True)

            df['type'] = df.apply(lambda x: 'one' if (x['right']-x['left']) >= x['image_width']/2.0 else 'two', axis=1)

            flag_one = False
            flag_two = False
            i_df = 0
            i_dict = 0
            tmp_list = list()
            result_dict = dict()
            while i_df < df.shape[0]:
                if df.type.iloc[i_df] == 'one' and flag_one == False and flag_two == True:
                    flag_one = True
                    flag_two = False
                    result_dict[i_dict] = tmp_list
                    tmp_list = list()
                    tmp_list.append(i_df)
                    i_dict += 1
                elif df.type.iloc[i_df] == 'two' and flag_one == True and flag_two == False:
                    flag_one = False
                    flag_two = True
                    result_dict[i_dict] = tmp_list
                    tmp_list = list()
                    tmp_list.append(i_df)
                    i_dict += 1
                elif df.type.iloc[i_df] == 'two' and flag_two == False:
                    flag_two = True
                    tmp_list.append(i_df)
                elif df.type.iloc[i_df] == 'two' and flag_two == True:
                    tmp_list.append(i_df)
                elif df.type.iloc[i_df] == 'one' and flag_one == False:
                    flag_one = True
                    tmp_list.append(i_df)
                elif df.type.iloc[i_df] == 'one' and flag_one == True:
                    tmp_list.append(i_df)
                i_df += 1

            result_dict[i_dict] = tmp_list

            result_df = pd.DataFrame()
            for i in range(len(result_dict)):
                tmp_df = df.iloc[result_dict[i], :]
                if 'two' == tmp_df['type'].iloc[0]:
                    tmp_df['marker'] = tmp_df.apply(lambda x: 'left' if x['left'] <= x['image_width']/2. else 'right', axis = 1)
                    tmp_df['right'] = tmp_df.right.apply(math.ceil)
                    
                    if tmp_df[tmp_df.marker == 'left'].shape[0] > 0:
                        min_l_left = math.floor(tmp_df[tmp_df.marker == 'left'].left.min())
                        #max_r_left = tmp_df[tmp_df.marker == 'left'].right.max()
                    
                    if tmp_df[tmp_df.marker == 'right'].shape[0] > 0:
                        min_l_right = math.floor(tmp_df[tmp_df.marker == 'right'].left.min())
                        #max_r_right = tmp_df[tmp_df.marker == 'right'].right.max()

                    result_tmp_df = pd.DataFrame()
                    
                    left_tmp_df = tmp_df[tmp_df.marker == 'left'].sort_values(by=['upper'], ascending=True)
                    if left_tmp_df.shape[0] > 0:
                        left_tmp_df.left = min_l_left
                    #left_tmp_df.right = max_r_left
                    
                    result_tmp_df = result_tmp_df.append(
                        left_tmp_df, ignore_index=True
                    )
                    
                    right_tmp_df = tmp_df[tmp_df.marker == 'right'].sort_values(by=['upper'], ascending=True)
                    if right_tmp_df.shape[0] > 0:
                        right_tmp_df.left = min_l_right
                    #right_tmp_df.right = max_r_right

                    result_tmp_df = result_tmp_df.append(
                        right_tmp_df, ignore_index=True
                    )
                    
                    tmp_df = result_tmp_df
                    
                result_df = result_df.append(tmp_df, ignore_index=True)
            
            #result_df.to_excel('show.xlsx')

            if 'marker' not in result_df.columns.tolist():
                result_df['marker'] = None

            result_df.marker = result_df.marker.fillna('one')

            bad_indexes = list()
            for i in range(result_df.shape[0]):
                for j in range(result_df.shape[0]):
                    bb1 = {
                        'x1': result_df.left.iloc[i],
                        'x2': result_df.right.iloc[i],
                        'y1': result_df.upper.iloc[i],
                        'y2': result_df.lower.iloc[i],
                        'marker': result_df.marker.iloc[i],
                        'label': result_df.label.iloc[i]
                    }
                    bb2 = {
                        'x1': result_df.left.iloc[j],
                        'x2': result_df.right.iloc[j],
                        'y1': result_df.upper.iloc[j],
                        'y2': result_df.lower.iloc[j],
                        'marker': result_df.marker.iloc[j],
                        'label': result_df.label.iloc[j]
                    }

                    # изменил условия на y1 и y2, + делаем проверку на маркеры

                    if bb1['y1'] >= bb2['y1'] and bb1['y2'] <= bb2['y2'] and (i!=j) and (bb1['label'] in ['text', 'title', 'list']) and (bb2['label'] in ['text', 'title', 'list']) and (bb1['marker'] == bb2['marker']):
                        bad_indexes.append(i)
                    elif (bb1['y2'] > bb2['y1']) and (i+1 == j) and (bb1['marker'] == bb2['marker']):
                        result_df.upper.iloc[j] = result_df.upper.iloc[i]
                        bad_indexes.append(i)

            result_df = result_df.drop(bad_indexes)
            
            for i in range(result_df.shape[0]-1):
                if result_df['type'].iloc[i] == 'one':
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]
                elif (result_df.marker.iloc[i] == 'left') and (result_df.marker.iloc[i+1] == 'left'):
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]
                elif ((result_df.marker.iloc[i] == 'right') and (result_df.marker.iloc[i+1] == 'right')) or ((result_df.marker.iloc[i] == 'right') and (result_df.marker.iloc[i+1] == 'one')):
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]
                # случай когда нет right вместе с left:
                elif (result_df.marker.iloc[i] == 'left') and (result_df.marker.iloc[i+1] == 'one'):
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]
                # самый сложный слуачй с left: поиск one после блока right или понимание это это самый нижний левый
                elif (result_df.marker.iloc[i] == 'left') and (result_df.marker.iloc[i+1] == 'right'):
                    k = copy.copy(i) + 1
                    while k <= result_df.shape[0]-1:
                        if result_df.marker.iloc[k] == 'one':
                            break
                        k += 1
                
                    # проверяем какая ситуация, если k < кол-ва данных, то значит нашли 
                    if k <= result_df.shape[0]-1:
                        result_df.lower.iloc[i] = result_df.upper.iloc[k]
                    else:
                        result_df.lower.iloc[i] += min((result_df.lower.iloc[i] - result_df.upper.iloc[i]) * 0.3, img.size[1] - result_df.lower.iloc[i])

            # прибавляем плюсом посленему элементу
            result_df.lower.iloc[result_df.shape[0]-1] += min((result_df.lower.iloc[result_df.shape[0]-1] - result_df.upper.iloc[result_df.shape[0]-1]) * 0.3, img.size[1] - result_df.lower.iloc[result_df.shape[0]-1])

            last_result_df = pd.DataFrame()
            last_result_df = last_result_df.append(
                result_df[(result_df['marker'] == 'one')|(result_df['marker'] == 'left')], ignore_index = True
            )
            last_result_df = last_result_df.append(
                result_df[(result_df['marker'] == 'right')], ignore_index = True
            )

            last_result_df = last_result_df.reset_index(drop=True)

            return last_result_df

        else:
            
            result_df = df.sort_values(by=['upper'], ascending=True)
            result_df = result_df.reset_index(drop=True)

            bad_indexes = list()
            for i in range(result_df.shape[0]):
                for j in range(result_df.shape[0]):
                    bb1 = {
                        'x1': result_df.left.iloc[i],
                        'x2': result_df.right.iloc[i],
                        'y1': result_df.upper.iloc[i],
                        'y2': result_df.lower.iloc[i],
                        'label': result_df.label.iloc[i]
                    }
                    bb2 = {
                        'x1': result_df.left.iloc[j],
                        'x2': result_df.right.iloc[j],
                        'y1': result_df.upper.iloc[j],
                        'y2': result_df.lower.iloc[j],
                        'label': result_df.label.iloc[j]
                    }

                    # bb2 входит в bb1

                    if bb1['y1'] >= bb2['y1'] and bb1['y2'] <= bb2['y2'] and (i!=j) and (bb1['label'] in ['text', 'title', 'list']) and (bb2['label'] in ['text', 'title', 'list']):
                        bad_indexes.append(i)

                    # iou, area_bb1, area_bb2 = self.calc_iou(bb1, bb2)
                    # if (iou > 0.0) and (i!=j) and (area_bb1 == 1.0):
                    #     bad_indexes.append(i)
            
            result_df = result_df.drop(bad_indexes)

            for i in range(result_df.shape[0]-1):
                if (result_df.upper.iloc[i+1] > result_df.upper.iloc[i]) and (result_df.label.iloc[i] in ['text', 'title', 'list']):
                    result_df.lower.iloc[i] = result_df.upper.iloc[i+1]

            return result_df


class PageElement:
    ''' Class for one element of the page '''
    def __init__(self, element_number: int, element_type: str, element_image: Image, bbox: tuple):
        self.element_number = element_number
        self.element_type = element_type
        self.element_image = element_image
        self.bbox = bbox


class TextPageElement(PageElement):
    def __init__(self, element_number: int, element_type: str, element_image: Image, bbox: tuple, langs: list, tessdata_dir: str, pdf_path: str, document_type: int, dpi: int, page_number: int):
        super().__init__(element_number, element_type, element_image, bbox)
        self.element_text = None
        self.langs = langs
        self.tessdata_dir = tessdata_dir
        self.pdf_path = pdf_path
        self.document_type = document_type
        self.dpi = dpi
        self.page_number = page_number
        self.extract_text()


    def extract_text(self):
        if self.document_type == 0 or self.document_type == 1:
            cropped_img = self.element_image
            half = 0.5
            cropped_img = cropped_img.resize( [int(half * s) for s in cropped_img.size] ) # conver image to 1/2 of start dpi
            use_lang = '+'.join(self.langs)
            custom_config = f'--oem 1 --psm 6 -l {use_lang} --tessdata-dir {self.tessdata_dir}'
            self.element_text = pytesseract.image_to_string(cropped_img, config=custom_config)
            self.element_text = str(self.element_text).replace('—', ' ')
            self.element_text = str(self.element_text).replace('\n', ' ')
        elif self.document_type == 2:
            bbox = str(int(self.bbox[0])), str(int(self.bbox[1])), str(int(self.bbox[2]) - int(self.bbox[0])), str(int(self.bbox[3]) - int(self.bbox[1]))
            page_num = str(self.page_number + 1)    
            args = ["/usr/bin/pdftotext",            
                    '-f', page_num, '-l', page_num,
                    '-r', str(self.dpi),
                    '-x', bbox[0], '-y',  bbox[1], '-W', bbox[2], '-H', bbox[3],
                    '-enc', 'UTF-8',
                    self.pdf_path, '-']
            res = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE) 
            text = res.stdout.decode('utf-8')    
            escape_char = re.compile(r'\\x[0123456789abcdef]+')
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
            text = re.sub('\n', ' ', text)
            if text == '' or text == ' ':
                cropped_img = self.element_image
                half = 0.5
                cropped_img = cropped_img.resize( [int(half * s) for s in cropped_img.size] ) # conver image to 1/2 of start dpi
                use_lang = '+'.join(self.langs)
                custom_config = f'--oem 1 --psm 6 -l {use_lang} --tessdata-dir {self.tessdata_dir}'
                text = pytesseract.image_to_string(cropped_img, config=custom_config)
                text = str(text).replace('—', ' ')
                text = str(text).replace('\n', ' ')
            self.element_text = text


class TablePageElement(PageElement):
    def __init__(self, element_number: int, element_type: str, element_image: Image, bbox: tuple):
        super().__init__(element_number, element_type, element_image, bbox)
        self.dataframe = None
    

    def save_dataframe_to_excel(self, path):
        if self.dataframe is not None:
            self.dataframe.to_excel(path)


class DocumentPage:
    ''' Class for page '''
    def __init__(self, page_number: int, path_to_page_image: str, layout_type: int, pdf_path: str, document_type: int, dpi: int):
        self.page_number = page_number
        self.path_to_page_image = path_to_page_image
        self.elements = list()
        self.layout_type = layout_type
        self.pdf_path = pdf_path
        self.document_type = document_type
        self.formulae = list()
        self.dpi = dpi


    def extract(self, segmentation_model: SegmentationModel, finding_formulas_model: FindingFormulasModel, langs: list, tessdata_dir: str):
        img = detectron2.data.detection_utils.read_image(self.path_to_page_image, format="RGB")
        df = segmentation_model.predict(img, self.layout_type)

        img = Image.fromarray(img)

        for i in range(df.shape[0]):
            cropped_img = img.crop((df.iloc[i,0], df.iloc[i,1], df.iloc[i,2], df.iloc[i,3]))
            if df.iloc[i,4] == 'title' or df.iloc[i,4] == 'text' or df.iloc[i,4] == 'list':
                tmp_page_elemnt = TextPageElement(element_number = i, element_type = df.iloc[i,4], element_image=cropped_img, bbox=(df.iloc[i,0], df.iloc[i,1], df.iloc[i,2], df.iloc[i,3]), langs=langs, tessdata_dir=tessdata_dir, pdf_path=self.pdf_path, document_type=self.document_type, dpi = self.dpi, page_number = self.page_number)
            elif df.iloc[i,4] == 'table':
                tmp_page_elemnt = TablePageElement(element_number = i, element_type = df.iloc[i,4], element_image=cropped_img, bbox=(df.iloc[i,0], df.iloc[i,1], df.iloc[i,2], df.iloc[i,3]))
            else:
                tmp_page_elemnt = PageElement(element_number = i, element_type = df.iloc[i,4], element_image=cropped_img, bbox=(df.iloc[i,0], df.iloc[i,1], df.iloc[i,2], df.iloc[i,3]))
            self.elements.append(tmp_page_elemnt)
        
        df_formulars = finding_formulas_model.predict(self.path_to_page_image)
        for i in range(df_formulars.shape[0]):
            cropped_img = img.crop((df_formulars.iloc[i,0], df_formulars.iloc[i,1], df_formulars.iloc[i,2], df_formulars.iloc[i,3]))
            tmp_page_elemnt = PageElement(element_number = i, element_type = 'formula', element_image=cropped_img, bbox=(df_formulars.iloc[i,0], df_formulars.iloc[i,1], df_formulars.iloc[i,2], df_formulars.iloc[i,3]))
            self.formulae.append(tmp_page_elemnt)

        
class Document:
    ''' Class for work with document '''
    def __init__(self, pdf_path: str, segmentation_model: SegmentationModel, finding_formulas_model: FindingFormulasModel, layout_type: int, langs: list, tessdata_dir: str, document_type: int, dpi: int):
        self.pdf_path = pdf_path
        self.pages = dict()
        self.segmentation_model = segmentation_model
        self.layout_type = layout_type
        self.languages = langs
        self.tessdata_dir = tessdata_dir
        self.document_type = document_type
        self.finding_formulas_model = finding_formulas_model
        self.dpi = dpi


    def pdf_to_jpg(self, file_path: str) -> dict:
        ''' Function for converting PDF to JPG using ImageMagick '''
        
        # Function for generating a random string
        rand_str = lambda n: ''.join([random.choice(string.ascii_lowercase) for i in range(n)])  
        
        rand_folder = rand_str(20)
        tmp_folder = f'tmp_{rand_folder}'
        os.system(f'mkdir ./tmp/{tmp_folder}')
        
        result = {
            'path_to_tmp_folder': './tmp/' + tmp_folder,
            'path_to_tmp_files': list()
        }
        
        image_path = result['path_to_tmp_folder'] + '/' + 'image.jpg'
        os.system(f'convert -density {self.dpi} -background white -alpha remove {file_path} -quality 100 {image_path}')
        
        files = os.listdir(result['path_to_tmp_folder'])
        
        for file in files:
            result['path_to_tmp_files'].append(result['path_to_tmp_folder'] + '/' + file)
        
        return result


    def add_table_to_docx(self, document, df):
        n_row, n_column = df.shape
        table = document.add_table(rows=1, cols=n_column)
        hdr_cells = table.rows[0].cells

        for i in range(n_column):
            if df.columns[i] != None:
                hdr_cells[i].text = str(df.columns[i])

        for i in range(n_row):
            row_cells = table.add_row().cells
            for j in range(n_column):
                if df.iloc[i, j] != None:
                    row_cells[j].text = str(df.iloc[i, j])
    
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return document  

    
    def build_docx_documnet(self, output_filename:str):
        document = DocumentDocx()
        num_pages = len(self.pages)
        for i in range(num_pages):
            page = self.pages[i]
            num_elements = len(page.elements)
            for j in range(num_elements):
                if page.elements[j].element_type == 'title' and len(page.elements[j].element_text) < 70:
                    head = document.add_heading(page.elements[j].element_text, level=1) 
                    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif page.elements[j].element_type == 'text' or page.elements[j].element_type == 'list' or (page.elements[j].element_type == 'title' and len(page.elements[j].element_text) >= 70):
                    text = page.elements[j].element_text
                    paragraph = document.add_paragraph(text)
                elif page.elements[j].element_type == 'figure':
                    document.add_picture(page.elements[j].image_path, width=Inches(5))
                elif page.elements[j].element_type == 'table':
                    if page.elements[j].dataframe is not None:
                        document = self.add_table_to_docx(document, page.elements[j].dataframe)
                    else:
                        document.add_picture(page.elements[j].image_path, width=Inches(5))
        
        document.save(f'./output/{self.output_dir}/{output_filename}')


    def pdfplumber_extractor(self, pdf_path: str):
        pdf = pdfplumber.open(pdf_path)    
        num_pages = len(pdf.pages)
        result_tables = pd.DataFrame(columns = ['page', 'left', 'upper', 'right', 'lower', 'data'])    
        for i in range(num_pages):       
            
            # Getting JPG page sizes
            jpg_path = self.pages[i].path_to_page_image
            im = cv2.imread(jpg_path)
            height_jpg, width_jpg, _ = im.shape  
                
            # Getting PDF page sizes
            p0 = pdf.pages[i] 
            width_plumber = float(p0.bbox[2] - p0.bbox[0])
            height_plumber = float(p0.bbox[3] - p0.bbox[1])
            
            # Extracting tables  
            bb = p0.find_tables()   
            tables = p0.extract_tables()       
            num_tables = len(tables)
            
            if num_tables > 0:     
                for j in range(num_tables):
                    df = pd.DataFrame(tables[j][1:], columns=tables[j][0])
                    
                    left_plumber = float(bb[j].bbox[0])
                    upper_plumber = float(bb[j].bbox[1])
                    right_plumber = float(bb[j].bbox[2])
                    lower_plumber = float(bb[j].bbox[3])
                    
                    left = left_plumber * width_jpg / width_plumber
                    upper = upper_plumber * height_jpg / height_plumber
                    right = right_plumber * width_jpg / width_plumber
                    lower = lower_plumber * height_jpg / height_plumber
                    
                    data = [{'page':i,'left':left,'upper':upper,'right':right,'lower':lower,'data': df}]                
                    result_tables.loc[len(result_tables.index)]=list(data[0].values()) 
                                    
        return result_tables # The function returns a dictionary of pages containing tables

    
    def calc_iou(self, bb1: dict, bb2: dict) -> float:
        """
        Calculate the Intersection over Union (IoU) of two bounding boxes.
        Parameters
        ----------
        bb1 : dict
            Keys: {'x1', 'x2', 'y1', 'y2'}
            The (x1, y1) position is at the top left corner,
            the (x2, y2) position is at the bottom right corner
        bb2 : dict
            Keys: {'x1', 'x2', 'y1', 'y2'}
            The (x, y) position is at the top left corner,
            the (x2, y2) position is at the bottom right corner
        Returns
        -------
        float
            in [0, 1]
        """
        assert bb1['x1'] < bb1['x2']
        assert bb1['y1'] < bb1['y2']
        assert bb2['x1'] < bb2['x2']
        assert bb2['y1'] < bb2['y2']

        # determine the coordinates of the intersection rectangle
        x_left = max(bb1['x1'], bb2['x1'])
        y_top = max(bb1['y1'], bb2['y1'])
        x_right = min(bb1['x2'], bb2['x2'])
        y_bottom = min(bb1['y2'], bb2['y2'])

        if x_right < x_left or y_bottom < y_top:
            return 0.0

        # The intersection of two axis-aligned bounding boxes is always an
        # axis-aligned bounding box
        intersection_area = (x_right - x_left) * (y_bottom - y_top)

        # compute the area of both AABBs
        bb1_area = (bb1['x2'] - bb1['x1']) * (bb1['y2'] - bb1['y1'])
        bb2_area = (bb2['x2'] - bb2['x1']) * (bb2['y2'] - bb2['y1'])

        # compute the intersection over union by taking the intersection
        # area and dividing it by the sum of prediction + ground-truth
        # areas - the interesection area
        iou = intersection_area / float(bb1_area + bb2_area - intersection_area)
        if iou < 0:
            return -1
        elif iou > 1:
            return -1
        else:
            return iou
    

    def match_pdfplumber_data(self, pdfplumber_data: pd.DataFrame):
        for i in range(pdfplumber_data.shape[0]):
            page_number = pdfplumber_data.page.iloc[i]
            for id_element, element in enumerate(self.pages[page_number].elements):
                bb1 = {
                    'x1': element.bbox[0],
                    'x2': element.bbox[2],
                    'y1': element.bbox[1],
                    'y2': element.bbox[3]
                }
                bb2 = {
                    'x1': pdfplumber_data.left.iloc[i],
                    'x2': pdfplumber_data.right.iloc[i],
                    'y1': pdfplumber_data.upper.iloc[i],
                    'y2': pdfplumber_data.lower.iloc[i]
                }
                iou = self.calc_iou(bb1, bb2)
                if iou >= 0.5:
                    self.pages[page_number].elements[id_element].dataframe = pdfplumber_data.data.iloc[i]


    def getting_output_dir(self):
        rand_str = lambda n: ''.join([random.choice(string.ascii_lowercase) for i in range(n)])
        rand_folder = rand_str(20)
        os.mkdir(f'./output/{rand_folder}')
        os.mkdir(f'./output/{rand_folder}/images/')
        os.mkdir(f'./output/{rand_folder}/tables/')
        os.mkdir(f'./output/{rand_folder}/formulae/')
        self.output_dir = rand_folder

    
    def save_nontext(self):
        for page_id in self.pages:
            for element_id, element in enumerate(self.pages[page_id].elements):
                if type(element).__name__ == 'PageElement':
                    element.image_path = f'./output/{self.output_dir}/images/page_{page_id}_element_number_{element_id}.jpg'
                    element.element_image.save(element.image_path,
                                              format= 'JPEG',
                                              quality = 100,
                                              icc_profile = element.element_image.info.get("icc_profile")
                                              )
                elif type(element).__name__ == 'TablePageElement':
                    if element.dataframe is not None:
                        element.table_path = f'./output/{self.output_dir}/tables/page_{page_id}_element_number_{element_id}.xlsx'
                        element.save_dataframe_to_excel(element.table_path)
                    else:
                        element.image_path = f'./output/{self.output_dir}/tables/page_{page_id}_element_number_{element_id}.jpg'
                        element.element_image.save(element.image_path,
                                              format= 'JPEG',
                                              quality = 100,
                                              icc_profile = element.element_image.info.get("icc_profile")
                                              )

            for element_id, element in enumerate(self.pages[page_id].formulae):
                if type(element).__name__ == 'PageElement':
                    element.image_path = f'./output/{self.output_dir}/formulae/page_{page_id}_element_number_{element_id}.jpg'
                    element.element_image.save(element.image_path,
                                              format= 'JPEG',
                                              quality = 100,
                                              icc_profile = element.element_image.info.get("icc_profile")
                                              )

        image_files = os.listdir(f'./output/{self.output_dir}/images/')
        if len(image_files) == 0:
            os.system(f'rm -r ./output/{self.output_dir}/images/')
        
        table_files = os.listdir(f'./output/{self.output_dir}/tables/')
        if len(table_files) == 0:
            os.system(f'rm -r ./output/{self.output_dir}/tables/')
        
        formula_files = os.listdir(f'./output/{self.output_dir}/formulae/')
        if len(formula_files) == 0:
            os.system(f'rm -r ./output/{self.output_dir}/formulae/')

    
    def create_zip(self) -> str:
        os.system(f'7z a ./output/{self.output_dir}.zip ./output/{self.output_dir}/*')
        return f'./output/{self.output_dir}.zip'

    
    def delete_tmp(self, tmp_path):
        os.system(f'rm -r {tmp_path}')
        os.system(f'rm -r ./output/{self.output_dir}')


    def convert(self, output_type: str, output_filename:str, to_zip: bool) -> str:

        # First step: convert to image
        imagemagick_output = self.pdf_to_jpg(self.pdf_path)


        # Second step: create pages:
        if len(imagemagick_output['path_to_tmp_files']) > 1:
            for i in range(len(imagemagick_output['path_to_tmp_files'])):
                tmp_page_number = int(re.findall(r'\d+', imagemagick_output['path_to_tmp_files'][i])[-1])
                tmp_page = DocumentPage(page_number = tmp_page_number, 
                                        path_to_page_image = imagemagick_output['path_to_tmp_files'][i], 
                                        layout_type = self.layout_type,
                                        pdf_path = self.pdf_path,
                                        document_type = self.document_type,
                                        dpi = self.dpi)
                self.pages[tmp_page_number] = tmp_page
        else:
            tmp_page_number = 0
            tmp_page = DocumentPage(page_number=tmp_page_number, 
                                    path_to_page_image=imagemagick_output['path_to_tmp_files'][0],
                                    layout_type=self.layout_type,
                                    pdf_path = self.pdf_path,
                                    document_type = self.document_type,
                                    dpi = self.dpi)
            self.pages[tmp_page_number] = tmp_page


        # Third stage: extracting all page elements
        for page_id in self.pages:
            self.pages[page_id].extract(
                segmentation_model = self.segmentation_model,
                langs = self.languages,
                tessdata_dir = self.tessdata_dir,
                finding_formulas_model = self.finding_formulas_model
            )

        
        # Fourth stage: extracting tables from a pdf document
        pdfplumber_data = self.pdfplumber_extractor(self.pdf_path)


        # Fifth step: comparison of pdf plumber_data
        self.match_pdfplumber_data(pdfplumber_data)

        # Sixth stage: Creating output folders
        self.getting_output_dir()

        # Seventh stage: The preservation of non-textual elements
        self.save_nontext()

        # Stage eight: creating a document of the desired format
        if output_type == 'docx':
            self.build_docx_documnet(output_filename)

        if to_zip:    
            # Stage nine: creating a zip archive
            path_to_zip = self.create_zip()
            # Step ten: deleting temporary files
            self.delete_tmp(imagemagick_output['path_to_tmp_folder'])
            return path_to_zip
        else:
            return self.output_dir
